#!/usr/bin/env python3
"""
Build the docxtpl-tagged finding template from a source docx.

Reads a hand-formatted source docx (whose visible text matches the layout
encoded in `PLACEHOLDER_MAP` below) and writes
`templates/finding_template.docx` with Jinja tags in place of the original
placeholder strings. The output is consumed by `skills/reporting.py` at
render time.

Re-run this script whenever the visual layout of the source template
changes. The layout itself (cell widths, fonts, borders, headers/footers)
is preserved — only the run text inside specific paragraphs/cells is
rewritten.

Usage:
    uv run python scripts/build_finding_template.py [SOURCE_DOCX] [OUTPUT_DOCX]

Defaults:
    SOURCE_DOCX = ~/Downloads/Finding Template.docx
    OUTPUT_DOCX = templates/finding_template.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = Path.home() / "Downloads" / "Finding Template.docx"
DEFAULT_OUTPUT = PROJECT_ROOT / "templates" / "finding_template.docx"


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace a paragraph's text while keeping the first run's formatting.

    Removes every direct child of the paragraph except the paragraph
    properties (`w:pPr`) and the first run (`w:r`). This also strips
    hyperlinks (`w:hyperlink`), which would otherwise leave orphaned link
    text behind. docxtpl scans run text for Jinja tags, so the surviving
    run carries the placeholder while keeping its inherited formatting.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    first_run = runs[0]._element
    parent = paragraph._element
    # Drop everything that isn't the paragraph props or the first run.
    for child in list(parent):
        tag = child.tag
        if tag == f"{_W_NS}pPr" or child is first_run:
            continue
        parent.remove(child)
    runs[0].text = new_text


def _replace_cell_text(cell, new_text: str) -> None:
    """Replace a table cell's text, collapsing to a single paragraph."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(new_text)
        return
    _replace_paragraph_text(paragraphs[0], new_text)
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def build(source: Path, output: Path) -> None:
    if not source.exists():
        raise FileNotFoundError(f"Source template not found: {source}")

    doc = Document(str(source))

    # --- Heading paragraph (outside tables) -------------------------------- #
    # P0: "BHI-OFFSEC-XX.XX.FXX - Relevant issue title"
    _replace_paragraph_text(doc.paragraphs[0], "{{ finding.id }} - {{ finding.title }}")

    # --- Table 0: Severity | … | CVSS 3.1 | x.x | Category | … ------------- #
    t0 = doc.tables[0]
    _replace_cell_text(t0.rows[0].cells[1], "{{ finding.severity }}")
    _replace_cell_text(t0.rows[0].cells[3], "{{ finding.cvss.score }}")
    _replace_cell_text(t0.rows[0].cells[5], "{{ finding.category }}")

    # --- Table 1: CVSS 3.1 Vector ----------------------------------------- #
    _replace_cell_text(doc.tables[1].rows[0].cells[1], "{{ finding.cvss.vector }}")

    # --- Table 2: Affected Asset(s) --------------------------------------- #
    _replace_cell_text(doc.tables[2].rows[0].cells[1], "{{ finding.affected }}")

    # --- Description / Impact / PoC / Remediation prose ------------------- #
    # The prose sits in paragraphs OUTSIDE the section-header tables.
    # After the heading-table rewrite, we re-resolve paragraph indices by
    # walking the body and locating prose paragraphs between the header
    # tables. To stay deterministic we hard-code the indices learned from
    # the source layout, then verify on each run.
    body_paragraphs = doc.paragraphs

    # Sanity checks — fail loudly if the source layout drifts.
    expected = {
        4: "A description of what was found",
        5: "The assessment uncovered an SQL injection",
        6: "highlight of the general impact",
        7: "Any relevant information to support the description",
        8: "set of concrete actions",
    }
    for idx, snippet in expected.items():
        if snippet not in body_paragraphs[idx].text:
            raise RuntimeError(
                f"Source template layout drift: paragraph {idx} no longer "
                f"contains expected snippet {snippet!r}. Re-inspect the "
                f"source docx and update build_finding_template.py."
            )

    # Description spans P4 + P5 in the source — collapse to one paragraph.
    _replace_paragraph_text(body_paragraphs[4], "{{ finding.description }}")
    _delete_paragraph(body_paragraphs[5])

    # Indices shift after the delete; re-resolve.
    body_paragraphs = doc.paragraphs
    _replace_paragraph_text(body_paragraphs[5], "{{ finding.impact }}")
    _replace_paragraph_text(body_paragraphs[6], "{{ finding.proof_of_concept }}")
    _replace_paragraph_text(body_paragraphs[7], "{{ finding.remediation }}")

    # --- References block (CVE / CWE / OWASP / MITRE / Blogs) ------------- #
    # Convert the 5 reference lines into a docxtpl paragraph loop. P9 (after
    # the deletion above) is the "References" header table's trailing blank;
    # the reference lines start at P10 in the source, P9 after the delete.
    body_paragraphs = doc.paragraphs
    ref_start = None
    for i, p in enumerate(body_paragraphs):
        if p.text.strip() == "CVE":
            ref_start = i
            break
    if ref_start is None:
        raise RuntimeError("Could not locate the CVE reference paragraph.")

    # P[ref_start]   → loop opener
    # P[ref_start+1] → loop body
    # P[ref_start+2] → loop closer
    # P[ref_start+3..] → delete (originally CWE/OWASP/MITRE/Blogs lines)
    _replace_paragraph_text(
        body_paragraphs[ref_start], "{%p for ref in finding.references %}"
    )
    _replace_paragraph_text(body_paragraphs[ref_start + 1], "{{ ref }}")
    _replace_paragraph_text(body_paragraphs[ref_start + 2], "{%p endfor %}")

    # Delete remaining reference placeholders (MITRE, Blogs).
    # We re-collect because index shifts after each delete.
    while True:
        body_paragraphs = doc.paragraphs
        leftover = None
        for p in body_paragraphs[ref_start + 3 :]:
            if p.text.strip() in {"MITRE", "Blogs and other specific references"}:
                leftover = p
                break
        if leftover is None:
            break
        _delete_paragraph(leftover)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"[+] Wrote {output}")


def main() -> int:
    src = Path(sys.argv[1]).expanduser() if len(sys.argv) > 1 else DEFAULT_SOURCE
    dst = Path(sys.argv[2]).expanduser() if len(sys.argv) > 2 else DEFAULT_OUTPUT
    print(f"[*] Source: {src}")
    print(f"[*] Output: {dst}")
    build(src, dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
